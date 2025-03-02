import json
import os
import re
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Параметры
KEY_W = 'usa'
PAGES_DIR = 'pars/pages/bt'
EV_JSON_PATH = 'ev.json'
IMPORTANT_CONTEXT_JSON = 'important_context.json'
ES_JSON_PATH = 'es.json'
OUTPUT_DOCX = f'Результаты анализа слова {KEY_W}.docx'

# Размеры окон
WINDOW_SIZE = 150  # Для подсчёта категорий
CONTEXT_WORDS = 50  # Для отображения контекста

# Загрузка данных
with open(EV_JSON_PATH, 'r', encoding="utf-8") as f:
    es = json.load(f)

with open(IMPORTANT_CONTEXT_JSON, 'r', encoding="utf-8") as f:
    kw = json.load(f)

with open(ES_JSON_PATH, 'r', encoding="utf-8") as f:
    es_categories = json.load(f)

# Загрузка цветовой карты
COLOR_MAP = {
    'v': RGBColor(0, 0, 255),  # синий
    'g': RGBColor(0, 128, 0),  # зелёный
    'r': RGBColor(255, 0, 0),  # красный
    'b': RGBColor(255, 255, 0),  # жёлтый
}

# Создание словаря: слово -> категория и цвет
word_category_map = {}
word_color_map = {}
for category, words in es_categories.items():
    prefix = category[0].lower()
    color = COLOR_MAP.get(prefix, None)
    for word in words:
        word_clean = word.strip().lower()
        if word_clean:
            word_category_map[word_clean] = category
            word_color_map[word_clean] = color


# Функция для добавления горизонтальной линии
def add_horizontal_line(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    # Создание элемента w:pBdr
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    # Создание нижней границы
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')  # Тип линии: сплошная
    bottom.set(qn('w:sz'), '6')  # Толщина линии (в десятых долей пункта)
    bottom.set(qn('w:space'), '1')  # Отступ между линией и текстом
    bottom.set(qn('w:color'), '000000')  # Цвет линии: черный

    # Добавление нижней границы к pBdr
    pBdr.append(bottom)

    # Добавление pBdr к pPr
    pPr.append(pBdr)


# Функция для выделения контекста в документе Word
def add_highlighted_context(document, words, index, total_words):
    start = max(index - CONTEXT_WORDS, 0)
    end = min(index + CONTEXT_WORDS + 1, total_words)
    context = words[start:end]
    paragraph = document.add_paragraph()

    for i, word in enumerate(context):
        word_lower = word.lower()
        run = paragraph.add_run(word + ' ')
        # Выделение ключевого слова
        if start + i == index:
            run.bold = True
            run.font.highlight_color = 7  # Желтый фон (используется стандартный код)
        elif word_lower in word_color_map:
            run.font.color.rgb = word_color_map[word_lower]

    # Добавление горизонтальной линии после контекста
    add_horizontal_line(document)


# Создание документа Word
doc = Document()

# Добавление заголовка
doc.add_heading('Результаты Анализа', level=1)

# Обработка файлов
files = os.listdir(PAGES_DIR)

for file in files:
    file_path = os.path.join(PAGES_DIR, file)
    with open(file_path, 'r', encoding="utf-8") as f:
        data = json.load(f)

    text = data.get('text', '').lower()
    date = data.get('date', 'Дата не указана')

    # Поиск ключевых слов
    found_positions = [m.start() for m in re.finditer(rf'\b{re.escape(KEY_W)}\b', text, re.IGNORECASE)]

    if found_positions:
        # Добавление даты
        doc.add_heading(f"Дата: {date}", level=2)

        # Разбивка текста на слова
        words = re.findall(r'\b\w+\b', text)
        total_words = len(words)

        # Создание обратного словаря для категорий
        keyword_counts = defaultdict(lambda: defaultdict(int))

        # Приведение ключевых слов к нижнему регистру
        kw_lower = [word.lower() for word in kw]

        for pos in found_positions:
            # Определение индекса слова
            word_index = len(re.findall(r'\b\w+\b', text[:pos]))
            if 0 <= word_index < total_words:
                # Определение окна для подсчёта
                start = max(word_index - WINDOW_SIZE, 0)
                end = min(word_index + WINDOW_SIZE + 1, total_words)
                context = words[start:end]

                for i, context_word in enumerate(context):
                    if (start + i) == word_index:
                        continue  # Пропуск ключевого слова
                    category = word_category_map.get(context_word)
                    if category:
                        keyword_counts[KEY_W][category] += 1

        # Добавление таблицы
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light List Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Группа'
        hdr_cells[1].text = 'Количество'

        counts = keyword_counts.get(KEY_W, {})
        if counts:
            for category, count in counts.items():
                row_cells = table.add_row().cells
                row_cells[0].text = category
                row_cells[1].text = str(count)
        else:
            row_cells = table.add_row().cells
            row_cells[0].text = "Нет связанных слов в контексте."
            row_cells[1].text = ""

        doc.add_paragraph()  # Пустая строка после таблицы

        # Вывод контекста для каждого вхождения
        for pos in found_positions:
            word_index = len(re.findall(r'\b\w+\b', text[:pos]))
            if 0 <= word_index < total_words:
                add_highlighted_context(doc, words, word_index, total_words)

# Сохранение документа
doc.save(OUTPUT_DOCX)
print(f"Результаты успешно сохранены в файл '{OUTPUT_DOCX}'.")
